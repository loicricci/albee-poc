"""
File processor for extracting text from various file formats.
Supports: PDF, Word docs, Excel, CSV, images, and URLs.
"""
import io
import base64
from typing import Dict, Any
from pathlib import Path
import requests
from bs4 import BeautifulSoup

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from PIL import Image
except ImportError:
    Image = None

from openai import OpenAI
import os

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file."""
    if PdfReader is None:
        raise ImportError("PyPDF2 is not installed. Install it with: pip install PyPDF2")
    
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                text_parts.append(f"[Page {page_num}]\n{text}")
        
        if not text_parts:
            return "[PDF contains no extractable text]"
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from Word document."""
    if DocxDocument is None:
        raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
    
    try:
        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        if not text_parts:
            return "[Document contains no extractable text]"
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from Word document: {str(e)}")


def extract_text_from_excel(file_content: bytes, filename: str) -> str:
    """Extract text from Excel file."""
    if pd is None:
        raise ImportError("pandas is not installed. Install it with: pip install pandas openpyxl")
    
    try:
        excel_file = io.BytesIO(file_content)
        
        # Read all sheets
        excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
        
        text_parts = []
        for sheet_name, df in excel_data.items():
            text_parts.append(f"[Sheet: {sheet_name}]")
            
            # Convert dataframe to string representation
            df_str = df.to_string(index=False, na_rep='')
            text_parts.append(df_str)
        
        if not text_parts:
            return "[Excel file contains no data]"
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from Excel file: {str(e)}")


def extract_text_from_csv(file_content: bytes) -> str:
    """Extract text from CSV file."""
    if pd is None:
        raise ImportError("pandas is not installed. Install it with: pip install pandas")
    
    try:
        csv_file = io.BytesIO(file_content)
        
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'iso-8859-1']:
            try:
                csv_file.seek(0)
                df = pd.read_csv(csv_file, encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Could not decode CSV file with common encodings")
        
        # Convert dataframe to string representation
        df_str = df.to_string(index=False, na_rep='')
        
        return f"[CSV Data]\n{df_str}"
    except Exception as e:
        raise ValueError(f"Failed to extract text from CSV file: {str(e)}")


def extract_text_from_markdown(file_content: bytes) -> str:
    """Extract text from Markdown file."""
    try:
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'iso-8859-1']:
            try:
                text = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Could not decode Markdown file with common encodings")
        
        if not text.strip():
            return "[Markdown file is empty]"
        
        return f"[Markdown Document]\n\n{text}"
    except Exception as e:
        raise ValueError(f"Failed to extract text from Markdown file: {str(e)}")


def process_image_with_gpt4_vision(file_content: bytes, filename: str) -> str:
    """Process image using GPT-4 Vision to extract text and describe content."""
    if Image is None:
        raise ImportError("Pillow is not installed. Install it with: pip install Pillow")
    
    try:
        # Validate it's an image
        img = Image.open(io.BytesIO(file_content))
        
        # Convert to base64
        base64_image = base64.b64encode(file_content).decode('utf-8')
        
        # Determine image format
        img_format = img.format.lower() if img.format else 'jpeg'
        mime_type = f"image/{img_format}"
        
        # Use GPT-4 Vision to analyze the image
        response = client.chat.completions.create(
            model="gpt-4o",  # GPT-4 with vision capabilities
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": """Please analyze this image and provide:
1. Any text visible in the image (transcribe it exactly)
2. A detailed description of what the image shows
3. Any relevant context or information that would be useful for understanding this image

Format your response as:
TEXT CONTENT:
[any visible text]

DESCRIPTION:
[detailed description]

CONTEXT:
[relevant context]"""
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=1000
        )
        
        extracted_info = response.choices[0].message.content
        return f"[Image: {filename}]\n\n{extracted_info}"
        
    except Exception as e:
        raise ValueError(f"Failed to process image: {str(e)}")


def process_file(file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
    """
    Process uploaded file and extract text content.
    
    Returns:
        Dict with 'text' and 'metadata' keys
    """
    file_ext = Path(filename).suffix.lower()
    
    try:
        # PDF files
        if file_ext == '.pdf' or 'pdf' in content_type.lower():
            text = extract_text_from_pdf(file_content)
            file_type = "PDF Document"
        
        # Word documents
        elif file_ext in ['.docx', '.doc'] or 'word' in content_type.lower():
            text = extract_text_from_docx(file_content)
            file_type = "Word Document"
        
        # Excel files
        elif file_ext in ['.xlsx', '.xls'] or 'spreadsheet' in content_type.lower() or 'excel' in content_type.lower():
            text = extract_text_from_excel(file_content, filename)
            file_type = "Excel Spreadsheet"
        
        # CSV files
        elif file_ext == '.csv' or 'csv' in content_type.lower():
            text = extract_text_from_csv(file_content)
            file_type = "CSV File"
        
        # Markdown files
        elif file_ext in ['.md', '.markdown'] or 'markdown' in content_type.lower():
            text = extract_text_from_markdown(file_content)
            file_type = "Markdown Document"
        
        # Image files
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'] or 'image' in content_type.lower():
            text = process_image_with_gpt4_vision(file_content, filename)
            file_type = "Image"
        
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        return {
            "text": text,
            "metadata": {
                "original_filename": filename,
                "file_type": file_type,
                "file_size": len(file_content),
                "content_type": content_type
            }
        }
    
    except Exception as e:
        raise ValueError(f"Error processing file '{filename}': {str(e)}")


def get_supported_extensions() -> list:
    """Return list of supported file extensions."""
    return [
        '.pdf',
        '.docx', '.doc',
        '.xlsx', '.xls',
        '.csv',
        '.md', '.markdown',
        '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'
    ]


def get_supported_mime_types() -> list:
    """Return list of supported MIME types."""
    return [
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel',
        'text/csv',
        'text/markdown',
        'text/x-markdown',
        'image/jpeg',
        'image/png',
        'image/gif',
        'image/bmp',
        'image/webp'
    ]


def extract_content_from_url(url: str) -> Dict[str, Any]:
    """
    Extract content from a URL.
    
    Returns:
        Dict with 'text' and 'metadata' keys
    """
    try:
        # Validate URL format
        if not url.startswith(('http://', 'https://')):
            raise ValueError("URL must start with http:// or https://")
        
        # Check for LinkedIn - requires special handling
        if 'linkedin.com' in url.lower():
            raise ValueError(
                "LinkedIn profiles require authentication and cannot be scraped directly. "
                "Please copy and paste the profile content manually, or export your LinkedIn data "
                "from Settings > Data Privacy > Get a copy of your data."
            )
        
        # Fetch the URL content
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        content_type = response.headers.get('Content-Type', '').lower()
        
        # Parse HTML content
        if 'text/html' in content_type or 'application/xhtml' in content_type:
            soup = BeautifulSoup(response.content, 'lxml')
            
            # Remove script and style elements
            for script in soup(['script', 'style', 'nav', 'footer', 'header']):
                script.decompose()
            
            # Get title
            title = soup.title.string if soup.title else 'No title'
            
            # Get meta description
            description = ''
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc and meta_desc.get('content'):
                description = meta_desc['content']
            
            # Get main content
            # Try to find main content area
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content') or soup.body
            
            if main_content:
                # Extract text
                text_parts = []
                
                # Add title
                if title != 'No title':
                    text_parts.append(f"Title: {title}")
                
                # Add description
                if description:
                    text_parts.append(f"Description: {description}")
                
                # Add main content
                paragraphs = main_content.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li'])
                for para in paragraphs:
                    text = para.get_text(strip=True)
                    if text and len(text) > 20:  # Filter out very short texts
                        text_parts.append(text)
                
                extracted_text = "\n\n".join(text_parts)
            else:
                # Fallback: get all text
                extracted_text = soup.get_text(separator='\n', strip=True)
            
            # Clean up excessive whitespace
            lines = [line.strip() for line in extracted_text.split('\n') if line.strip()]
            extracted_text = '\n'.join(lines)
            
            if not extracted_text or len(extracted_text) < 50:
                raise ValueError("Could not extract meaningful content from URL")
            
            return {
                "text": f"[URL: {url}]\n[Title: {title}]\n\n{extracted_text}",
                "metadata": {
                    "source_url": url,
                    "title": title,
                    "description": description,
                    "content_type": content_type,
                    "file_type": "Web Page"
                }
            }
        
        # Handle PDF from URL
        elif 'application/pdf' in content_type:
            text = extract_text_from_pdf(response.content)
            return {
                "text": f"[URL: {url}]\n\n{text}",
                "metadata": {
                    "source_url": url,
                    "content_type": content_type,
                    "file_type": "PDF from URL"
                }
            }
        
        # Handle plain text
        elif 'text/plain' in content_type:
            return {
                "text": f"[URL: {url}]\n\n{response.text}",
                "metadata": {
                    "source_url": url,
                    "content_type": content_type,
                    "file_type": "Text File from URL"
                }
            }
        
        else:
            raise ValueError(f"Unsupported content type from URL: {content_type}")
    
    except requests.exceptions.Timeout:
        raise ValueError(f"Request timeout while fetching URL (30s limit). The website may be slow or unreachable.")
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 403:
            raise ValueError(f"Access forbidden (403). The website is blocking automated requests. Try copying the content manually.")
        elif e.response.status_code == 404:
            raise ValueError(f"Page not found (404). Please check the URL is correct.")
        elif e.response.status_code == 429:
            raise ValueError(f"Too many requests (429). The website is rate-limiting. Please try again later.")
        else:
            raise ValueError(f"HTTP error {e.response.status_code}: {str(e)}")
    except requests.exceptions.ConnectionError:
        raise ValueError(f"Connection error. Please check your internet connection and that the URL is accessible.")
    except requests.exceptions.RequestException as e:
        raise ValueError(f"Failed to fetch URL: {str(e)}")
    except Exception as e:
        raise ValueError(f"Error processing URL '{url}': {str(e)}")

